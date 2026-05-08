from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "submission" / "seongnam_10min_access_visualization_result_official_data.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="D7D1C5"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_text(cell, text, bold=False, size=9.5, color="182B38"):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.line_spacing = 1.35
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)


def add_body(paragraph, text, bold=False, size=10.2, color="182B38"):
    paragraph.paragraph_format.line_spacing = 1.45
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)


def add_section_title(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(47, 111, 104)


def add_bullets(cell, items, size=9.6):
    cell.text = ""
    for item in items:
        paragraph = cell.add_paragraph(style=None)
        paragraph.paragraph_format.left_indent = Cm(0.15)
        paragraph.paragraph_format.first_line_indent = Cm(-0.15)
        paragraph.paragraph_format.line_spacing = 1.35
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run("· " + item)
        run.font.name = "맑은 고딕"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor(24, 43, 56)
    if cell.paragraphs and not cell.paragraphs[0].text:
        cell._tc.remove(cell.paragraphs[0]._p)


def add_form_row(table, label, content, label_fill="F2F0EA"):
    row = table.add_row()
    label_cell, content_cell = row.cells
    set_cell_shading(label_cell, label_fill)
    set_cell_text(label_cell, label, bold=True, size=9.3, color="253544")
    label_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    content_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_border(label_cell)
    set_cell_border(content_cell)
    if isinstance(content, list):
        add_bullets(content_cell, content)
    else:
        set_cell_text(content_cell, content, size=9.8)
    return row


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.25)
    section.bottom_margin = Cm(1.05)
    section.left_margin = Cm(1.55)
    section.right_margin = Cm(1.55)

    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    style.font.size = Pt(10)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("성남 10분 생활필수 접근권 격차 지도")
    run.bold = True
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(17)
    run.font.color.rgb = RGBColor(24, 43, 56)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_body(
        sub,
        "공식 공개 데이터 기반 정책 시각화 결과서",
        size=10.5,
        color="5B6B6A",
    )

    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(3.25)
    table.columns[1].width = Cm(13.7)

    add_form_row(table, "작품 제목", "성남 10분 생활필수 접근권 격차 지도")
    add_form_row(
        table,
        "주제 선택 배경",
        "생활서비스가 행정구역 안에 있어도 고령층·1인가구가 실제로 쉽게 닿지 못하면 체감 접근권은 낮아진다. 성남의 의료, 약국, 돌봄, 교통 기반시설을 함께 보아 정책 우선순위를 설명하고자 했다.",
    )

    add_form_row(
        table,
        "사용 데이터",
        [
            "성남시 의료기관 현황(data.go.kr): 대표 동별 의료 접근성 점수 산정",
            "성남시 약국 현황(data.go.kr): 대표 동별 약국 접근성 점수 산정",
            "성남시 노인종합복지관 현황(data.go.kr) + 복지시설 좌표 PDF(seongnam.go.kr): 돌봄 거점 거리 산정",
            "경기교통정보센터 버스정류소 집계 + OpenStreetMap 정류장 좌표: 교통 접근성 산정",
            "행정동별 인구·세대 및 1인세대 현황(data.go.kr): 고령층, 1인가구, 지원수요 지표 산정",
        ],
    )

    add_form_row(
        table,
        "데이터 활용 내용",
        [
            "공식 CSV를 내려받아 대표 생활권 10개 동으로 매핑하고, 의료기관·약국 수를 접근성 점수로 표준화했다.",
            "인구, 고령층, 1인세대, 고령 1인세대 비율을 결합해 지원수요 지표를 만들었다.",
            "노인복지관은 공식 시설 목록과 성남시 좌표 PDF를 결합했으며, 교통은 실제 버스정류장 좌표와의 거리로 표시했다.",
        ],
    )
    add_form_row(
        table,
        "시각화 방법",
        [
            "MapLibre 지도 위에 대표 생활권과 서비스 거점을 배치하고, 점수 구간을 색상과 크기로 구분했다.",
            "서비스별 small multiple 지도, 우선순위 산점도, 정책 브리프 카드를 연결해 10초 안에 핵심을 읽도록 구성했다.",
            "v2에서는 줄간격, 여백, 카드 내부 간격을 확대해 발표자료 캡처로 바로 사용할 수 있게 다듬었다.",
        ],
    )
    add_form_row(
        table,
        "시각화 결과물 설명",
        [
            "첫 화면은 생활서비스가 있는 것과 10분 안에 닿는 것은 다르다는 메시지와 대표 지도를 함께 보여준다.",
            "두 번째 화면은 의료, 약국, 교통, 돌봄별로 부족 지역이 다르게 나타난다는 점을 비교한다.",
            "세 번째 화면은 접근이 어렵고 지원수요가 높은 생활권을 우선 점검 대상으로 드러낸다.",
            "네 번째 화면은 의료 보완, 이동 접근, 돌봄, 복합 보완 유형별 실행 과제를 정책 카드로 정리한다.",
        ],
    )
    add_form_row(
        table,
        "활용방안 및 기대효과",
        "보건소, 교통, 복지 부서가 같은 화면에서 우선 점검 생활권을 공유할 수 있다. 향후 보행 네트워크, 통신·카드 집계 데이터가 추가되면 실제 10분 도달성과 이용 수요를 결합한 정책 의사결정 도구로 확장할 수 있다.",
    )

    add_section_title(doc, "출처 URL")
    source_table = doc.add_table(rows=0, cols=2)
    source_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    source_table.autofit = False
    sources = [
        ("공모전 공고", "https://seongnam.go.kr/city/1000052/30001/bbsView.do?idx=381551"),
        ("의료기관", "https://www.data.go.kr/data/15000890/fileData.do"),
        ("약국", "https://www.data.go.kr/data/15000848/fileData.do"),
        ("노인복지관", "https://www.data.go.kr/data/15000901/fileData.do"),
        ("버스정류소 집계", "https://gits.gg.go.kr/gtdb/web/trafficDb/publicTransport/busStop.do"),
        ("버스정류장 좌표", "https://overpass-api.de/"),
        ("인구·세대", "https://www.data.go.kr/data/15007386/fileData.do"),
        ("1인세대", "https://www.data.go.kr/data/15061108/fileData.do"),
        ("복지시설 좌표", "https://www.seongnam.go.kr/contents/down/10458_6.pdf"),
    ]
    for label, url in sources:
        cells = source_table.add_row().cells
        set_cell_text(cells[0], label, bold=True, size=7.9, color="2F6F68")
        set_cell_text(cells[1], url, size=7.6, color="3E4D56")
        set_cell_border(cells[0], "E5E0D6")
        set_cell_border(cells[1], "E5E0D6")

    doc.add_page_break()
    add_section_title(doc, "데이터 해석상 주의 및 확장 계획")
    notes = [
        "현재 결과서는 공식 공개 데이터를 앱 구조에 반영한 최신본이다. 의료, 약국, 인구, 1인세대, 노인복지관은 실제 공개 파일을 내려받아 재산출했다.",
        "교통 지표는 실제 버스정류장 좌표와 대표 동 중심점의 직선거리 기반이다. 발표 시 보행 네트워크 이동시간은 아직 반영하지 않았다고 설명한다.",
        "공모전 제공 민간 통신·카드 데이터는 제공 경로와 활용 가능성을 확인했지만, 현재 앱 CSV에는 직접 반영하지 않았다. 원자료를 확보하면 체류·소비·방문 강도 보정 지표로 추가한다.",
        "정책 적용 단계에서는 보행 네트워크, 경사, 횡단 대기, 야간·휴일 운영 여부를 추가해 실제 10분 도달성 지표로 확장한다.",
    ]
    for item in notes:
        paragraph = doc.add_paragraph()
        add_body(paragraph, "· " + item, size=10.2, color="3E4D56")

    add_section_title(doc, "재현 방법")
    for item in [
        "원천 파일은 `data/raw_official/`에 보관했다.",
        "`python scripts/prepare-official-data.py`를 실행하면 앱용 `data/accessibility_metrics.csv`와 `data/service_points.csv`가 다시 생성된다.",
        "같은 컬럼 구조를 유지하면 지도, 서비스별 카드, 산점도, 정책 브리프 카드가 자동으로 갱신된다.",
    ]:
        paragraph = doc.add_paragraph()
        add_body(paragraph, "· " + item, size=10.2, color="3E4D56")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
